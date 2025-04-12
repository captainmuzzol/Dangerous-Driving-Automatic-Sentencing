# coding=utf-8
# import PyMuPDF
# from typing import final
import datetime
import os
import re
import shutil
import subprocess

import docx
import fitz
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

import ExtraDef
import ScbgScan

'''字典'''
# 实刑条例
shixing_dict = {
    1: "致使他人轻伤",
    2: "有酒驾或醉驾前科",
    3: "无证驾驶机动车",
    4: "驾驶中型以上货车"
}
# 数字转化
num_dict = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9}
unit_dict = {'十': 10, '百': 100, '千': 1000, '万': 10000}

'''初值'''
global fact_start
global fact_end, lxjy_C
global N_middle_word
N_middle_word = ""
fact_start = 0
fact_end = 0
i = 0
j = 0
chengshu = 0.75  # 乘数初值
lxqk_d = 0  # 量刑日期初值
lxqk_m = 0  # 量刑月份初值
minus_date = datetime.date.today()  # 日期差初值
wj_minus_date = 30000  # 危驾日期差初值

'''函数'''


# 数字阿拉伯转中文
def atc(arab_num):
    digits = {'0': '', '1': '一', '2': '二', '3': "三", '4': '四', '5': '五', '6': '六', '7': '七', '8': '八', '9': '九'}
    units = ['', '十', '百', '千', '万']
    result = ''
    try:
        l = len(arab_num) - 1
        m = 0
        need_zero = False
        while l >= 0:
            if arab_num[l] != '0':
                result = digits[arab_num[l]] + units[m] + result
            else:
                # 如果这一位是零，则需要在数字前面加上一个零
                if not need_zero and result:
                    result = '零' + result
                    need_zero = True
            l -= 1
            m += 1
            # 防止过万
            if l == 5 or m == 5:
                result = units[m] + result
                m = 0
                need_zero = False
    except:
        print("报错：数字阿拉伯转中文失败")
    return result


# 阿拉伯转中文数字
def cta(chinese_num):
    result = 0
    num = 0
    for k in chinese_num:
        if k in num_dict:
            num = num_dict[k]
        elif k in unit_dict:
            result += num * unit_dict[k]
            num = 0
        elif k.isdigit():
            num = num * 10 + int(k)
        else:
            pass
    result += num
    result = str(result)
    return result


'''开始量刑'''


def SC(file):
    global fact, J_middle_word, final_result, info, extra, lxjy, lxjy_C, lxqk, chengshu
    global lxjg, N_middle_word, douhao, fact_start, fact_end, qianke, minus_date
    chengshu = 0.75
    minus_date = datetime.date.today()  # 日期差初值
    wj_minus_date = 30000  # 危驾日期差初值
    lxjy = ""  # 量刑建议初值
    lxjy_C = ""
    info = ""  # 嫌疑人基本信息
    extra = ""  # 额外情节识别
    J_middle_word = 0  # 酒精含量初值
    N_middle_word = ""  # 嫌疑人姓名初值
    final_result = ""  # 第一次量刑结果
    fact = ""
    fact_start = 0
    fact_end = 0
    qianke = ''

    # 获取工作路径，以获得检察官名字
    file_load = os.getcwd()
    inquisitor_file = file_load + '/模板/inquisitor.txt'
    with open(inquisitor_file, 'r', errors='ignore') as file_2:
        inquisitor = file_2.read()
    if inquisitor == "通用":
        inquisitor = "XXX"
    if inquisitor == "选择检察官":
        inquisitor = "XXX"
    print('Pdfget调试信息：已识别到检察官为：' + inquisitor)

    # 读取输出文件并拆分成自然段
    with open(file, 'r', errors='ignore') as file:
        text = file.read()

        # 将文本按换行符和句号拆分成自然段
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        paragraphs = [p.strip() for para in paragraphs for p in para.split('.') if p.strip()]

        # 获取息与事实段落
        for para in paragraphs:
            if para.find("经依法侦查查明") != -1:
                fact_start = paragraphs.index(para)
            if para.find("认定上述") != -1:
                fact_end = paragraphs.index(para)
        i = fact_start
        while i <= fact_end - 1:
            # fact_list.append(paragraphs[i])
            fact = fact + paragraphs[i].strip('\n')
            i += 1
        j = 0
        while j <= fact_start - 1:
            info = info + paragraphs[j].strip('\n')
            j += 1

    # 查找酒精含量
    J_start_keyword = '量为'
    J_end_keyword = 'mg'
    J_middle_words = []
    # for para in paragraphs:
    if J_start_keyword in fact and J_end_keyword in fact:
        start_index = fact.find(J_start_keyword) + len(J_start_keyword)
        end_index = fact.find(J_end_keyword)
        if start_index < end_index:
            try:
                J_middle_text = fact[start_index:end_index].strip()
                # 提取两个关键词中间的词语
                J_middle_word = J_middle_text.split()[0]
                J_middle_words.append(J_middle_word)
            except BaseException:
                print("未知错误")
        else:
            pass
    else:
        pass
    # print("info:"+info)
    # 查赵嫌疑人姓名
    N_start_keyword = '犯罪嫌疑人'
    N_end_keyword = '男，'  # Windows
    # N_end_keyword = '涉嫌危险驾驶罪一案'  # Mac
    N_middle_words = []
    if N_start_keyword in info and N_end_keyword in info:
        # print("1!")
        start_index = info.find(N_start_keyword) + len(N_start_keyword)
        end_index = info.find(N_end_keyword)
        if start_index < end_index:
            # print("2!")
            N_middle_text = info[start_index:end_index].strip()
            # 提取两个关键词中间的词语
            N_middle_word = N_middle_text.split()[0]
            # print('调试信息：N_middle_word' + N_middle_word+'\n')
            N_middle_words.append(N_middle_word)
        else:
            pass
    else:
        N_end_keyword = '女，'
        if N_start_keyword in info and N_end_keyword in info:
            start_index = info.find(N_start_keyword) + len(N_start_keyword)
            end_index = info.find(N_end_keyword)
            if start_index < end_index:
                N_middle_text = info[start_index:end_index].strip()
                # 提取两个关键词中间的词语
                N_middle_word = N_middle_text.split()[0]  # 有时候会出现多个姓名，这里只取第一个
                # print('调试信息：N_middle_word' + N_middle_word+'\n')
                N_middle_words.append(N_middle_word)
    # 不男不女
    if N_middle_word == "":
        N_start_keyword = '犯罪嫌疑人'
        N_end_keyword = '男性，'  # Windows
        # N_end_keyword = '涉嫌危险驾驶罪一案'  # Mac
        N_middle_words = []
        if N_start_keyword in info and N_end_keyword in info:
            # print("1!")
            start_index = info.find(N_start_keyword) + len(N_start_keyword)
            end_index = info.find(N_end_keyword)
            if start_index < end_index:
                # print("2!")
                N_middle_text = info[start_index:end_index].strip()
                # 提取两个关键词中间的词语
                N_middle_word = N_middle_text.split()[0]
                # print('调试信息：N_middle_word' + N_middle_word+'\n')
                N_middle_words.append(N_middle_word)
            else:
                pass
        else:
            N_end_keyword = '女性，'
            if N_start_keyword in info and N_end_keyword in info:
                start_index = info.find(N_start_keyword) + len(N_start_keyword)
                end_index = info.find(N_end_keyword)
                if start_index < end_index:
                    N_middle_text = info[start_index:end_index].strip()
                    # 提取两个关键词中间的词语
                    N_middle_word = N_middle_text.split()[0]  # 有时候会出现多个姓名，这里只取第一个
                    # print('调试信息：N_middle_word' + N_middle_word+'\n')
                    N_middle_words.append(N_middle_word)

    N_middle_word = N_middle_word.rstrip("，")
    # 输出结果
    douhao = "," + "，"
    final_result += '犯罪嫌疑人: ' + N_middle_word + '，' + '\n' + '酒精含量为: ' + str(J_middle_word) + 'mg/100ml'

    NumGet = int(J_middle_word)
    Basic_num = "%.2f" % (NumGet / 80)
    # 量刑计算部分
    lxqk = 0  # 量刑情节（初步量刑）初值
    lxjg = 0  # 量刑结果初值
    fj = 0  # 罚金初值
    ewfjpd = 0  # 额外罚金判断初值
    shixing_attention = 0  # 实刑提示
    hx = 0  # 缓刑期初值
    moto_attention = 0  # 摩托初值
    money_attention = 0  # 损失量初值
    czqj = ""
    judge = 0

    # 前科劣迹判断
    info = ExtraDef.Rep_info(info)  # 修改info表述格式
    # print("info:" + info)
    global info_split_list
    info_split_list = (re.split('[。；]', info))  # 此处还需要split分号，分割信息部分(re.split('。|；', info))
    for info_split in info_split_list:
        if (info_split.startswith("2")) or info_split.startswith("1"):
            qianke = qianke + info_split + '；'
            try:  # 防止公安不写“日”
                qk_date = info_split[:info_split.index("日")] + "日"
                first_date = datetime.date(int(qk_date[:info_split.index("年")]),
                                           int(cta(qk_date[info_split.index("年") + 1:info_split.index("月")])),
                                           int(cta(qk_date[
                                                   info_split.index("月") + 1:info_split.index("日")])))  # 转化为时间格式
            except:
                qk_date = info_split[:info_split.index("因")] + "日"
                first_date = datetime.date(int(qk_date[:info_split.index("年")]),
                                           int(cta(qk_date[info_split.index("年") + 1:info_split.index("月")])),
                                           int(cta(qk_date[
                                                   info_split.index("月") + 1:info_split.index("因")])))  # 转化为时间格式
            now_date = datetime.date.today()  # 获取当前时间
            minus_date = now_date - first_date  # 计算日期差
            minus_date = minus_date.days  # 转化为int类型，不然无法比较大小
            # print(minus_date)

            if info_split.find("醉") != -1 or info_split.find("饮") != -1 or info_split.find("危险驾驶") != -1:
                wj_minus_date = minus_date
                if 1800 >= minus_date >= 60 and (info_split.find("醉") != -1):
                    czqj = czqj + "g"  # 5年醉驾
                    judge = "1"
                elif 1080 >= minus_date >= 60 and info_split.find("饮") != -1:
                    czqj = czqj + "g"  # 3年醉驾
                    judge = "1"
                elif 1080 <= minus_date and info_split.find("饮") != -1:
                    judge = "1"
                    czqj = czqj + "r"  # 有关前科劣迹
                elif 1800 <= minus_date:
                    judge = "1"
                    czqj = czqj + "r"  # 有关前科劣迹
            elif minus_date >= 60:
                judge = "1"
                czqj = czqj + "q"  # 无关前科劣迹
    print('Pdfget调试信息|前科劣迹：', qianke)
    # print(czqj)

    '''数据传递 额外情节识别'''
    if fact.find("事故") != -1 and fact.find("两车") != -1:  # 撞车事故
        judge = "1"
        czqj = czqj + "f"
    if fact.find("交通事故") != -1 and fact.find("发生碰撞") != -1:  # 撞车事故
        judge = "1"
        czqj = czqj + "f"
    if fact.find("受伤") != -1 and not fact.find("轻伤") != -1:  # 一般受伤
        czqj = czqj + "d"
    if fact.find("摩托") != -1 or fact.find("二轮") != -1:  # 摩托车    暂无法使用
        judge = "1"
        czqj = czqj + "4"
        moto_attention = 1
        final_result += '\n' + "对不起，摩托车相关案件暂时无法量刑，请关闭程序。"
        exit()
    if fact.find("无有效机动车") != -1 or fact.find("无证") != -1 or fact.find("扣留期间") != -1:  # 无证驾驶
        judge = "1"
        czqj = czqj + "h"
    if fact.find("牌照") != -1 or fact.find("无牌") != -1:  # 无牌驾驶
        judge = "1"
        czqj = czqj + "m"
    if fact.find("轻伤") != -1 or fact.find("轻伤") != -1:  # 轻伤
        judge = "1"
        czqj = czqj + "m"
    if fact.find("营运") != -1 or fact.find("营运") != -1:  # 营运
        judge = "1"
        czqj = czqj + "m"
    if fact.find("驾车逃离") != -1 or fact.find("逃离") != -1:  # 事故后逃离现场
        judge = "1"
        czqj = czqj + "c"

    # 如果只有事故没有人受伤，且酒精含量小于170，可不起诉
    czqj_cant_do = ['q', 'f', 'qf']
    if czqj.rstrip('0123456789') in czqj_cant_do and int(J_middle_word) < 170 and not fact.find("伤") != -1:
        judge = "0"

    jishu = float(Basic_num)
    '''量刑规范'''
    if jishu * 0.75 <= 0.8:
        final_result += '\n' + "酒精含量低于最小值，请确认是否正确。"
    if jishu * 0.75 == jishu * 0.75 <= 1.2:
        lxqk_m = 1
        lxqk_d = 0
    elif jishu * 0.75 <= 1.4:
        lxqk_m = 1
        lxqk_d = 10
    elif jishu * 0.75 <= 1.6:
        lxqk_m = 1
        lxqk_d = 15
    elif jishu * 0.75 <= 1.8:
        lxqk_m = 1
        lxqk_d = 20
    elif jishu * 0.75 <= 2.2:
        lxqk_m = 2
        lxqk_d = 0
    elif jishu * 0.75 <= 2.4:
        lxqk_m = 2
        lxqk_d = 10
    elif jishu * 0.75 <= 2.6:
        lxqk_m = 2
        lxqk_d = 15
    elif jishu * 0.75 <= 2.8:
        lxqk_m = 2
        lxqk_d = 20
    elif jishu * 0.75 <= 3.2:
        lxqk_m = 3
        lxqk_d = 0
    elif jishu * 0.75 <= 3.4:
        lxqk_m = 3
        lxqk_d = 10
    elif jishu * 0.75 <= 3.6:
        lxqk_m = 3
        lxqk_d = 15
    elif jishu * 0.75 <= 3.8:
        lxqk_m = 3
        lxqk_d = 20
    elif jishu * 0.75 <= 4.2:
        lxqk_m = 4
        lxqk_d = 0
    elif jishu * 0.75 <= 4.4:
        lxqk_m = 4
        lxqk_d = 10
    elif jishu * 0.75 <= 4.6:
        lxqk_m = 4
        lxqk_d = 15

    '''缓刑期计算'''
    if jishu * 0.75 <= 1.4:
        hx = 2
    elif jishu * 0.75 < 2:
        hx = 3
    elif 2 <= jishu * 0.75 <= 2.6:
        hx = 4
    elif 2.6 < jishu * 0.75 <= 3.2:
        hx = 5
    elif 3.2 < jishu * 0.75 <= 3.8:
        hx = 6

    if czqj == "":  # 设卡（无额外情节）
        final_result += '\n' + "未识别到任何额外情节。"
        if jishu * 0.75 <= 1.593:
            lxjy = "相对不起诉"
            lxjy_C = "相对不起诉"
            final_result += '\n' + '\n' + '    ▶>>>建议量刑：可考虑做不起诉处理<<<◀' + '\n'
        else:
            fj = int(fj) + lxqk_m * 2000
            fj = int(fj)
            if lxqk_d > 0:
                fj = fj + (lxqk_d - 5) / 5 * 500
                fj = int(fj)
            lxjy = str(lxqk_m) + '个月' + str(lxqk_d) + '天，缓刑' + str(hx) + '个月，并处罚金人民币' + str(fj) + '元'
            lxjy_C = atc(str(lxqk_m)) + '个月' + atc(str(lxqk_d)) + '天，缓刑' + atc(str(hx)) + '个月，并处罚金人民币' + atc(
                str(fj)) + '元'
            final_result += '\n' + '\n' + '     ▶>>>建议量刑：' + str(lxqk_m) + '个月' + str(lxqk_d) + '天，缓刑' + str(
                hx) + '个月，并处罚金人民币' + str(fj) + '元<<<◀'
            final_result += '\n' + '本结果仅供参考，请结合实际情况量刑'
    else:  # 非设卡（有额外情节）
        hx = 0  # 避免反复加缓刑
        fj = 0
        # czqj = "0"    #检验用
        if czqj.find('a') != -1:  # 自撞
            chengshu += 0.05
            lxjg = int(lxjg) + 5
            money_attention = 1  # 判断损失量
            extra += "自撞；"
        if czqj.find('b') != -1:  # 自撞后逃离现场
            chengshu += 0.05
            lxjg = int(lxjg) + 5
            extra += "自撞后逃离；"
        if czqj.find('c') != -1:  # 造成他人受伤或损失后逃逸/轻微抗拒执法
            chengshu += 0.1
            lxjg = int(lxjg) + 10
            extra += "造成他人受伤或损失后逃逸/轻微抗拒执法；"
        if czqj.find('d') != -1:  # 对他人造成一般伤势
            lxjg = int(lxjg) + 15
            extra += "对他人造成一般伤势；"
        if czqj.find('e') != -1:  # 造成他人轻伤
            jishu += 1
            lxjg = int(lxjg) + 20
            shixing_attention = 1
            extra += "造成他人轻伤；"
        if czqj.find('f') != -1:  # 损坏他人/公共财物
            chengshu += 0.1
            lxjg = int(lxjg) + 15
            money_attention = 1  # 判断损失量
            extra += "损坏他人财物；"
        if czqj.find('g') != -1:  # 有三年酒驾、五年醉驾情节
            jishu += 1
            lxjg = int(lxjg) + 20
            shixing_attention = 2
            extra += "有三年酒驾、五年醉驾情节；"
        if czqj.find('h') != -1:  # 无证驾驶
            # print(wj_minus_date)
            if wj_minus_date < 1800:
                jishu += 0.5
            else:
                jishu += 1
            lxjg = int(lxjg) + 20
            shixing_attention = 3
            extra += "无证驾驶；"
        if czqj.find('i') != -1:  # 校车业务/营运车辆（旅客运输、危险化学品运输）
            jishu += 1
            lxjg = int(lxjg) + 20
            ewfjpd = int(ewfjpd) + 1
            extra += "校车业务/营运车辆（旅客运输、危险化学品运输）；"
        if czqj.find('j') != -1:  # 营运车辆（货物运输，不含旅客）
            jishu += 1
            lxjg = int(lxjg) + 20
            extra += "营运车辆（货物运输，不含旅客）；"
        if czqj.find('k') != -1:  # 取保候审期间再次醉酒驾驶
            jishu += 2
            lxjg = int(lxjg) + 40
            extra += "取保候审期间再次醉酒驾驶；"
        if czqj.find('m') != -1:  # 无牌上路
            jishu += 1
            lxjg = int(lxjg) + 20
            extra += "无牌上路；"
        if czqj.find('n') != -1:  # 中型以上机动车(和营运车辆不重复选择）
            jishu += 2
            lxjg = int(lxjg) + 40
            shixing_attention = 4
            extra += "中型以上机动车；"
        if czqj.find('o') != -1:  # 在高速公路上行驶
            jishu += 1
            lxjg = int(lxjg) + 20
            extra += "高速公路行驶；"
        if czqj.find('p') != -1:  # 严重超员、超载、超速驾驶
            jishu += 1
            lxjg = int(lxjg) + 20
            extra += "严重超员、超载、超速驾驶；"
        if czqj.find('q') != -1:  # 与危驾无关的前科劣迹
            chengshu += 0.05
            lxjg = int(lxjg) + 5
            extra += "与危驾无关的前科劣迹；"
        if czqj.find('r') != -1:  # 与危驾有关的前科劣迹
            chengshu += 0.05
            lxjg = int(lxjg) + 5
            extra += "与危驾有关的前科劣迹；"
        if czqj.find('1') != -1:  # 自首（明知他人报警不逃离现场）
            chengshu -= 0.05
            lxjg = int(lxjg) - 5
            extra += "自首；"
        if czqj.find('2') != -1:  # 取得谅解未赔偿
            chengshu -= 0.05
            lxjg = int(lxjg) - 5
            extra += "取得谅解未赔偿；"
        if czqj.find('3') != -1:  # 赔偿并取得谅解
            chengshu -= 0.1
            lxjg = int(lxjg) - 10
            extra += "赔偿并取得谅解；"
        if czqj.find('4') != -1:  # 驾驶摩托车
            chengshu -= 0.3
            lxjg = int(lxjg) - 30
            moto_attention = 1  # 判断摩托提示
            extra += "驾驶摩托车；"
        # 避免乘数增减过高过低
        if chengshu <= 0.6:
            chengshu = 0.6
        elif chengshu >= 1:
            chengshu = 1

        # 刑期计算
        # print(str(jishu) + '\n' + str(chengshu))
        jishu = jishu * chengshu
        jishu = round(jishu, 1)

        '''量刑规范'''
        if jishu <= 0.8:
            final_result += '\n' + '该酒精含量低于最低量刑标准'
        if jishu == jishu <= 1.2:
            lxqk_m = 1
            lxqk_d = 0
        elif jishu <= 1.4:
            lxqk_m = 1
            lxqk_d = 10
        elif jishu <= 1.6:
            lxqk_m = 1
            lxqk_d = 15
        elif jishu <= 1.8:
            lxqk_m = 1
            lxqk_d = 20
        elif jishu <= 2.2:
            lxqk_m = 2
            lxqk_d = 0
        elif jishu <= 2.4:
            lxqk_m = 2
            lxqk_d = 10
        elif jishu <= 2.6:
            lxqk_m = 2
            lxqk_d = 15
        elif jishu <= 2.8:
            lxqk_m = 2
            lxqk_d = 20
        elif jishu <= 3.2:
            lxqk_m = 3
            lxqk_d = 0
        elif jishu <= 3.4:
            lxqk_m = 3
            lxqk_d = 10
        elif jishu <= 3.6:
            lxqk_m = 3
            lxqk_d = 15
        elif jishu <= 3.8:
            lxqk_m = 3
            lxqk_d = 20
        elif jishu <= 4.2:
            lxqk_m = 4
            lxqk_d = 0
        elif jishu <= 4.4:
            lxqk_m = 4
            lxqk_d = 10
        elif jishu <= 4.6:
            lxqk_m = 4
            lxqk_d = 15
        elif jishu > 4.6:
            final_result += '\n' + '超出最大值'

        '''缓刑期计算'''
        if jishu <= 1.4:
            hx = hx + 2
        elif jishu < 2:
            hx = hx + 3
        elif 2 <= jishu <= 2.6:
            hx = hx + 4
        elif 2.6 < jishu <= 3.2:
            hx = hx + 5

        if extra != "":
            final_result += '\n' + ("已识别到的额外情节有：" + extra)
        else:
            final_result += '\n' + ("未识别到任何额外情节。")

        '''完善月份日期显示'''
        # lxqk_d = int(lxqk_d) + lxjg   # 加法计算刑期

        if lxqk_d < 0:
            lxqk_m = int(lxqk_m) - 1
            lxqk_d = 30 + lxqk_d
        elif 60 > lxqk_d >= 30:
            lxqk_d = lxqk_d - 30
            lxqk_m = int(lxqk_m) + 1
        elif 90 > lxqk_d >= 60:
            lxqk_d = lxqk_d - 60
            lxqk_m = int(lxqk_m) + 2
        elif 120 > lxqk_d >= 90:
            lxqk_d = lxqk_d - 90
            lxqk_m = int(lxqk_m) + 3
        elif 150 > lxqk_d >= 120:
            lxqk_d = lxqk_d - 120
            lxqk_m = int(lxqk_m) + 4
        elif 180 > lxqk_d >= 150:
            lxqk_d = lxqk_d - 150
            lxqk_m = int(lxqk_m) + 5
        elif 210 > lxqk_d >= 180:
            lxqk_d = lxqk_d - 180
            lxqk_m = int(lxqk_m) + 6

        # 去除5、25天的量刑情况
        extra_no_num = extra.rstrip('0123456789')  # 去除减刑项
        if lxqk_d == 5:
            if len(extra_no_num) >= 2:
                lxqk_d += 5
            else:
                lxqk_d -= 5
        elif lxqk_d == 25:
            if len(extra_no_num) >= 2:
                lxqk_d += 5
                lxqk_m += 1
            else:
                lxqk_d -= 5

        '''罚金计算'''
        if ewfjpd == 0:  # 没有额外罚金的情况
            fj = int(fj) + lxqk_m * 2000
            if lxqk_d > 0:
                fj = fj + (lxqk_d - 5) / 5 * 500
        else:  # 有额外罚金的情况（营运等）
            fj = int(fj) + lxqk_m * 5000
            if lxqk_d > 0:
                fj = fj + (lxqk_d - 5) / 5 * 1250
        fj = int(fj)
        '''额外缓刑计算'''
        if lxqk_d >= 20 and lxqk_m * 2 >= hx:
            hx = hx + 1
        if lxqk_m == 2 and lxqk_d == 0 and jishu < 2.75:
            hx = 3
        # 如果前科过多，则增加缓刑期限
        if czqj.count("q") >= 3 and lxqk_m == 2 and hx == 3:
            hx += 1

        '''量刑输出'''
        if shixing_attention == 0:  # 没有上述的实刑条件时
            if int(judge) == 0:
                lxjy = "相对不起诉"
                lxjy_C = "相对不起诉"
                final_result += '\n' + '\n' + '     ▶>>>建议量刑：可考虑做不起诉处理<<<◀' + '\n'
            else:
                fj = int(fj)
                lxjy = str(lxqk_m) + '个月' + str(lxqk_d) + '天，缓刑' + str(hx) + '个月，并处罚金人民币' + str(fj) + '元'
                lxjy_C = atc(str(lxqk_m)) + '个月' + atc(str(lxqk_d)) + '天，缓刑' + atc(str(hx)) + '个月，并处罚金人民币' + atc(
                    str(fj)) + '元'
                final_result += '\n' + '\n' + '     ▶>>>建议量刑：' + str(lxqk_m) + '个月' + str(lxqk_d) + '天，缓刑' + str(
                    hx) + '个月，并处罚金人民币' + str(fj) + '元<<<◀''\n'
        else:  # 包含上述的实刑条件时
            fj = int(fj)
            lxjy = str(lxqk_m) + '个月' + str(lxqk_d) + '天，并处罚金人民币' + str(fj) + '元'
            lxjy_C = atc(str(lxqk_m)) + '个月' + atc(str(lxqk_d)) + '天，缓刑' + atc(str(hx)) + '个月，并处罚金人民币' + atc(
                str(fj)) + '元'
            final_result += '\n' + '\n' + '      ▶>>>建议量刑：' + str(lxqk_m) + '个月' + str(lxqk_d) + '天' + ',并处罚金人民币' + str(
                fj) + '元<<<◀''\n'
            final_result += '\n' + '注意：因其' + shixing_dict[shixing_attention] + '，建议不做缓刑处理'
        if money_attention == 1:  # 判断损失量
            final_result += '\n' + ' 如果造成较大金额损失（如>十万元），可适当增加刑期。如果已赔偿谅解，可适当减少刑期''\n'
        if moto_attention == 1:
            final_result += '\n' + '如果是驾驶摩托且酒精含量<200、认罪悔罪时，可考虑不起诉处理，其他情况也可考虑缓刑'

        final_result += '\n' + ' 起诉意见书获取信息有限！本结果仅供参考，请结合实际情况量刑（更新时间:2023.3.22)'

        # 标记
    '''生成需要的书'''

    i = 0
    info_new = ""

    # 修改量刑后调整
    LX_change = ''
    LX_file = file_load + '/模板/LX_change.txt'
    try:
        with open(LX_file, 'r', errors='ignore') as LX_file2:
            LX_change = LX_file2.read()
    except:
        pass

    if LX_change != '':
        lxjy = LX_change
    else:
        pass
    print('lxjy', lxjy)
    if lxjy == "建议量刑：可考虑做不起诉处理":
        pass
    else:
        LX_change_month = lxjy[5]
        print("LX_change_month", LX_change_month)
        LX_change_day = lxjy[(lxjy.find("月") + 1):(lxjy.find("天"))]
        print("LX_change_day", LX_change_day)
        LX_change_fj = lxjy[(lxjy.find("币") + 1):(lxjy.find("元"))]
        print("LX_change_fj", LX_change_fj)
        if lxjy.find("缓刑") != -1:
            LX_change_hx = lxjy[(lxjy.find("缓刑") + 2):(lxjy.find("月，") - 1)]
            print("LX_change_hx", LX_change_hx)
            lxjy_C = atc(LX_change_month) + '个月' + atc(LX_change_day) + '天，缓刑' + atc(LX_change_hx) + '个月，并处罚金人民币' + \
                     atc(str(LX_change_fj)) + '元'
        else:
            print('Pdfget调试信息:20240611临时检查2：',LX_change_month, LX_change_day,LX_change_fj)
            lxjy_C = atc(LX_change_month) + '个月' + atc(LX_change_day) + '天，' + '并处罚金人民币' + \
                     atc(str(LX_change_fj)) + '元'
            print('Pdfget调试信息:20240611临时检查1')
    
    print('lxjy', lxjy)
    print('lxjy_C', lxjy_C)
    
    # 对lxjy_C做调整
    if lxjy_C.find("一十") != -1:
        lxjy_C = lxjy_C.replace("一十", "十")
    else:
        pass
    if lxjy_C.find("月天") != -1:
        lxjy_C = lxjy_C.replace("月天", "月")
    else:
        pass
    lxjy_C = lxjy_C + "。"  # 补上句号
    # print(lxjy_C)
    if lxjy_C != "相对不起诉。":
        lxjy_C = "拘役" + lxjy_C
    else:
        pass

    # 信息截取
    print('Pdfget调试信息N_middle_word:', N_middle_word)
    R_name = N_middle_word.rstrip(douhao)
    print('Pdfget调试信息R_name:', R_name)
    for info_split in info_split_list:
        # print(info_split_list[i]+"\n")
        if not (info_split.startswith("2") or info_split.startswith("因") or info_split.startswith("该")):
            info_new = info_new + info_split_list[i]
            i = i + 1
        else:
            break
    info_new = "姓名：" + info_new.lstrip("犯罪嫌疑人") + "。"

    # 加入下划线的函数
    def add_underline_to_text(doc, search_text):
        for para in doc.paragraphs:
            for run in para.runs:
                if search_text in run.text:
                    run.font.underline = True
                    # run.font.underline_color = WD_UNDERLINE.SINGLE

    # 加粗的函数
    def change_text(path, word, font_name, font_size, bold, color, align):
        wordfile = Document(path)
        for paragraph in wordfile.paragraphs:
            for run in paragraph.runs:
                if word in run.text:
                    run.font.size = Pt(font_size)  # 字号
                    run.font.bold = bold  # 加粗
                    run.font.color.rgb = color  # 颜色
                    if align != "":
                        paragraph.alignment = align  # 对齐方式WD_PARAGRAPH_ALIGNMENT.CENTER
                    else:
                        pass
                    run.font.name = font_name  # 字体
                    r = run._element.rPr.rFonts
                    r.set(qn('w:eastAsia'), font_name)
                else:
                    pass
        wordfile.save(path)

    # 对起诉书的基础信息重新裁切
    info_new2 = ""
    i = 0
    info_split_list[0] = info_split_list[0].lstrip("犯罪嫌疑人")
    for info_split in info_split_list:
        # print(info_split_list[i]+"\n")
        if not info_split.startswith("犯罪嫌疑人"):
            info_new2 = info_new2 + info_split_list[i] + "。"
            # info_new2 = info_new2.lstrip("")
            i = i + 1
        else:
            break
    info_new2 = info_new2.replace(',', '，')  # 替换中文逗号

    '''信息再提取'''
    # 提取身份证号码
    id_number = ExtraDef.id_get(info_new2)

    # 提取性别
    sex = '未成功提取性别'
    sex_judge = int(id_number[-2])
    if sex_judge % 2 == 0:
        sex = '女'
    else:
        sex = '男'

    # 提取年龄
    birth_year = int(id_number[6:10])
    birth_month = int(id_number[10:12])
    birth_day = int(id_number[12:14])
    birthday = datetime.date(birth_year, birth_month, birth_day)
    today = datetime.date.today()
    age = today.year - birthday.year - ((today.month, today.day) < (birthday.month, birthday.day))
    age = str(age) + '岁'

    # 提取户籍
    print('Pdfget调试信息info：', info)
    place = ''
    place_now = ''
    place_get = ExtraDef.place_get(info)
    if place_get.find("户籍所在地") != -1:
        if place_get.find("现住") != -1:
            place_now_index = place_get.find('现住')
            place = place_get[6:place_now_index]
            place_now = place_get[(place_now_index + 2):]
        else:
            place = place_get[6:]
            place_now = ''
    else:
        place = place_get
    print('Pdfget调试信息place:', place)
    print('Pdfget调试信息place_now:', place_now)

    # 具结书
    def JJS_find_replace(doc_filename, find_text, replace_text):
        # 打开 Word 文档
        doc = docx.Document(doc_filename)

        # 遍历文档中的段落
        for p in doc.paragraphs:
            if re.search(find_text, p.text):
                # 使用 sub 方法替换文本
                p.text = re.sub(find_text, replace_text, p.text)

        # 遍历文档中的表格
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if re.search(find_text, p.text):
                            # 使用 sub 方法替换文本
                            p.text = re.sub(find_text, replace_text, p.text)

        # 保存新文件
        new_filename = "认罪认罚具结书" + "（" + R_name + "）.docx"
        doc.save(new_filename)

    # 起诉书
    def QSS_find_replace(doc_filename, find_text, replace_text):
        # 打开 Word 文档
        doc = docx.Document(doc_filename)

        # 遍历文档中的段落
        for p in doc.paragraphs:
            if re.search(find_text, p.text):
                # 使用 sub 方法替换文本
                p.text = re.sub(find_text, replace_text, p.text)

        # 遍历文档中的表格
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if re.search(find_text, p.text):
                            # 使用 sub 方法替换文本
                            p.text = re.sub(find_text, replace_text, p.text)

        # 保存新文件
        new_filename = "起诉书" + "（" + R_name + "）.docx"
        doc.save(new_filename)

    # 不起诉决定书
    def BSS_find_replace(doc_filename, find_text, replace_text):
        # 打开 Word 文档
        doc = docx.Document(doc_filename)

        # 遍历文档中的段落
        for p in doc.paragraphs:
            if re.search(find_text, p.text):
                # 使用 sub 方法替换文本
                p.text = re.sub(find_text, replace_text, p.text)

        # 遍历文档中的表格
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if re.search(find_text, p.text):
                            # 使用 sub 方法替换文本
                            p.text = re.sub(find_text, replace_text, p.text)

        # 保存新文件
        new_filename = "不起诉决定书" + "（" + R_name + "）.docx"
        doc.save(new_filename)

    # 落实三个规定报告表
    def SGGD_find_replace(doc_filename, find_text, replace_text):
        # 打开 Word 文档
        doc = docx.Document(doc_filename)

        # 遍历文档中的段落
        for p in doc.paragraphs:
            if re.search(find_text, p.text):
                # 使用 sub 方法替换文本
                p.text = re.sub(find_text, replace_text, p.text)

        # 遍历文档中的表格
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if re.search(find_text, p.text):
                            # 使用 sub 方法替换文本
                            p.text = re.sub(find_text, replace_text, p.text)

        # 保存新文件
        new_filename = "落实“三个规定”报告表" + "（" + R_name + "）.docx"
        doc.save(new_filename)

    # 公诉意见书
    def GSYJS_find_replace(doc_filename, find_text, replace_text):
        # 打开 Word 文档
        doc = docx.Document(doc_filename)

        # 遍历文档中的段落
        for p in doc.paragraphs:
            if re.search(find_text, p.text):
                # 使用 sub 方法替换文本
                p.text = re.sub(find_text, replace_text, p.text)

        # 遍历文档中的表格
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if re.search(find_text, p.text):
                            # 使用 sub 方法替换文本
                            p.text = re.sub(find_text, replace_text, p.text)

        # 保存新文件
        new_filename = "公诉意见书" + "（" + R_name + "）.docx"
        doc.save(new_filename)

    # 讯问笔录
    def XWBL_find_replace(doc_filename, find_text, replace_text):
        # 打开 Word 文档
        doc = docx.Document(doc_filename)

        # 遍历文档中的段落
        for p in doc.paragraphs:
            if re.search(find_text, p.text):
                # 使用 sub 方法替换文本
                p.text = re.sub(find_text, replace_text, p.text)

        # 遍历文档中的表格
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if re.search(find_text, p.text):
                            # 使用 sub 方法替换文本
                            p.text = re.sub(find_text, replace_text, p.text)

        # 保存新文件
        new_filename = "讯问笔录" + "（" + R_name + "）.docx"
        doc.save(new_filename)

    # 审查报告
    def SCBG_find_replace(doc_filename, find_text, replace_text):
        # 打开 Word 文档
        doc = docx.Document(doc_filename)

        # 遍历文档中的段落
        for p in doc.paragraphs:
            if re.search(find_text, p.text):
                # 使用 sub 方法替换文本
                p.text = re.sub(find_text, replace_text, p.text)

        # 遍历文档中的表格
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if re.search(find_text, p.text):
                            # 使用 sub 方法替换文本
                            p.text = re.sub(find_text, replace_text, p.text)

        # 保存新文件
        new_filename = "审查报告" + "（" + R_name + "）.docx"
        doc.save(new_filename)

    # 获取桌面地址
    def get_desktop():
        return os.path.join(os.path.expanduser("~"), 'Desktop')

    desktop = get_desktop()  # 判断是否存在文件夹，若不存在，则创建
    # print(desktop)
    print("PDFget调试信息（R_name:",R_name)
    desktop_path = desktop + '/' + R_name + "危险驾驶案"
    if not os.path.exists(desktop_path):
        os.mkdir(desktop_path)
    else:
        pass

    # 字体设置函数
    def modify_specific_word(doc_path, word, font_size, font_name):
        document = Document(doc_path)
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                if word in run.text:
                    font = run.font
                    font.bold = True
                    font.size = Pt(font_size)
                    font.name = font_name
        document.save(doc_path)

    if qianke == '':
        qianke = '（无前科）'
    else:
        qianke = qianke.rstrip('；')
        qianke += '。'
    # 具结书 生成
    JJS = "认罪认罚具结书" + "（" + R_name + "）.docx"
    JJS_find_replace(r"模板/认罪认罚具结书.docx", "{基础信息}", info_new2)  # 基础信息
    JJS_find_replace(JJS, "{姓名}", R_name)  # 姓名
    JJS_find_replace(JJS, "{犯罪事实}", fact.lstrip("经依法侦查查明："))  # 犯罪事实
    JJS_find_replace(JJS, "{量刑建议}", lxjy_C)  # 量刑建议
    JJS_find_replace(JJS, "执勤民警", "温岭市公安局交警大队民警")  # 民警替换
    JJS_find_replace(JJS, "经温岭市公安司法鉴定中心鉴定", "经抽血检验")  # 经……抽血检验
    if lxjy_C == '相对不起诉。':
        JJS_find_replace("认罪认罚具结书" + "（" + R_name + "）.docx", "4.本案适用速裁程序。", "")  # 删除速裁程序
    doc = Document(JJS)
    search_text_list = [info_new2, '指控本人犯罪事实', lxjy_C]  # 下划线
    for search_text in search_text_list:
        add_underline_to_text(doc, search_text)
    doc.save(JJS)

    # 讯问笔录 生成
    XWBL = "讯问笔录" + "（" + R_name + "）.docx"
    XWBL_find_replace(r"模板/讯问笔录.docx", "{基础信息}", info_new)  # 基础信息
    XWBL_find_replace(XWBL, "{姓名}", R_name)  # 姓名
    XWBL_find_replace(XWBL, "{强制措施}", ExtraDef.catch(info))  # 强制措施
    XWBL_find_replace(XWBL, "{事实经过}", ExtraDef.human_fact(fact, R_name))  # 自述事实
    XWBL_find_replace(XWBL, "{归案时间}", ExtraDef.time_fact(fact))  # {归案时间}
    print('Pdfget调试信息:走到这里（笔录成功生成）')
    XWBL_find_replace(XWBL, "{归案地点}", ExtraDef.place_fact(fact) + "时")  # {归案地点}、
    XWBL_find_replace(XWBL, "{前科}", qianke)  # 前科劣迹
    XWBL_find_replace(XWBL, "XXX", inquisitor)  # 检察官名字替换
    XWBL_find_replace(XWBL, "执勤民警", "温岭市公安局交警大队民警")  # 民警替换
    XWBL_find_replace(XWBL, "经温岭市公安司法鉴定中心鉴定", "经抽血检验")  # 经……抽血检验
    doc = Document(XWBL)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    doc.styles['Normal'].font.size = Pt(16)
    search_text_list = [info_new, R_name, '讯问人', qianke, ExtraDef.catch(info), ExtraDef.human_fact(fact, R_name)
                        , ExtraDef.time_fact(fact), ]  # 下划线
    for search_text in search_text_list:
        add_underline_to_text(doc, search_text)
    doc.save(XWBL)

    # print('Pdfget调试信息：lxjy_C', lxjy_C)
    if lxjy_C == "相对不起诉。":
        print('Pdfget调试信息：生成不起诉决定书', lxjy_C)
        # 不起诉决定书 生成
        BSS = "不起诉决定书" + "（" + R_name + "）.docx"
        BSS_find_replace(r"模板/不起诉决定书.docx", "{基础信息}", info_new2)  # 基础信息
        BSS_find_replace(BSS, "{姓名}", R_name)  # 姓名
        BSS_find_replace(BSS, "{犯罪事实}", fact.lstrip("经依法侦查查明："))  # 犯罪事实
        BSS_find_replace(BSS, "{量刑建议}", lxjy_C)  # 量刑建议
        BSS_find_replace(BSS, "我局", "温岭市公安局")  # 我局替换为温岭市公安局
        BSS_find_replace(BSS, "执勤民警", "温岭市公安局交警大队民警")  # 民警替换
        BSS_find_replace(BSS, "经温岭市公安司法鉴定中心鉴定", "经抽血检验")  # 经……抽血检验
        doc = Document(BSS)
        doc.save(BSS)
    else:
        # 起诉书 生成
        QSS = "起诉书" + "（" + R_name + "）.docx"
        QSS_find_replace(r"模板/起诉书.docx", "{基础信息}", info_new2)  # 基础信息
        QSS_find_replace(QSS, "{姓名}", R_name)  # 姓名
        QSS_find_replace(QSS, "{犯罪事实}", fact.lstrip("经依法侦查查明："))  # 犯罪事实
        QSS_find_replace(QSS, "{量刑建议}", lxjy_C)  # 量刑建议
        QSS_find_replace(QSS, "我局", "温岭市公安局")  # 我局替换为温岭市公安局
        QSS_find_replace(QSS, "XXX", inquisitor)  # 检察官名字替换
        QSS_find_replace(QSS, "执勤民警", "温岭市公安局交警大队民警")  # 民警替换
        QSS_find_replace(QSS, "经温岭市公安司法鉴定中心鉴定", "经抽血检验")  # 经……抽血检验
        doc = Document(QSS)
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
        doc.styles['Normal'].font.size = Pt(16)
        doc.save(QSS)

        # 公诉意见书 生成
        GSYJS_find_replace(r"模板/公诉意见书.docx", "{姓名}", R_name)  # 姓名
        GSYJS_find_replace("公诉意见书" + "（" + R_name + "）.docx", "{量刑建议}", lxjy_C)  # 量刑建议
        doc = Document("公诉意见书（" + R_name + "）.docx")
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
        doc.styles['Normal'].font.size = Pt(16)
        doc.save("公诉意见书（" + R_name + "）.docx")

    # 三个规定
    if lxjy == "相对不起诉":
        SGGD_find_replace(r"模板/落实“三个规定”报告表.docx", "{姓名}", R_name)  # 姓名
        SGGD_find_replace("落实“三个规定”报告表" + "（" + R_name + "）.docx", "{量刑建议}", lxjy_C)  # 量刑建议
        doc = Document("落实“三个规定”报告表" + "（" + R_name + "）.docx")
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        doc.styles['Normal'].font.size = Pt(14)
        doc.save("落实“三个规定”报告表" + "（" + R_name + "）.docx")
        # modify_specific_word(doc, '实行贯彻落实“三个规定”一案一填报', 10, '宋体')
    else:
        SGGD_find_replace(r"模板/落实“三个规定”报告表.docx", "{姓名}", R_name)  # 姓名
        SGGD_find_replace("落实“三个规定”报告表" + "（" + R_name + "）.docx", "{量刑建议}", "起诉")  # 量刑建议
        doc = Document("落实“三个规定”报告表" + "（" + R_name + "）.docx")
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        doc.styles['Normal'].font.size = Pt(14)
        doc.save("落实“三个规定”报告表" + "（" + R_name + "）.docx")
        # modify_specific_word(doc, '实行贯彻落实“三个规定”一案一填报', 10, '宋体')

    # 审查报告  生成
    SCBG = "审查报告" + "（" + R_name + "）.docx"
    pages = ScbgScan.get_bookmarks(desktop_path + '/诉讼证据卷.pdf')  # 获取页码


    if inquisitor != "丁林":  # 非丁林版审查报告
        SCBG_find_replace(r"模板/审查报告.docx", "{基础信息}", info_new)  # 基础信息
        SCBG_find_replace(SCBG, "{姓名}", R_name)  # 姓名
        SCBG_find_replace(SCBG, "{犯罪事实}", fact.lstrip("经依法侦查查明："))  # 犯罪事实
        SCBG_find_replace(SCBG, "{量刑建议}", lxjy_C)  # 量刑建议
        SCBG_find_replace(SCBG, "{事实经过}", ExtraDef.human_fact(fact, R_name))  # 自述事实
        SCBG_find_replace(SCBG, "{归案时间}", ExtraDef.time_fact(fact))  # {归案时间}
        SCBG_find_replace(SCBG, "{归案地点}", ExtraDef.place_fact(fact) + "时")  # {归案地点}
        SCBG_find_replace(SCBG, "{前科}", qianke)  # 前科劣迹
        SCBG_find_replace(SCBG, "{户籍}", place)  # 户籍
        SCBG_find_replace(SCBG, "{现住}", '现住：' + place_now)  # 现住
        SCBG_find_replace(SCBG, "{性别}", sex)  # 性别
        SCBG_find_replace(SCBG, "{年龄}", age)  # 年龄
        SCBG_find_replace(SCBG, "{酒精含量}", J_middle_word + 'mg/100ml')  # 酒精含量
        SCBG_find_replace(SCBG, "{身份证号码}", id_number)  # 身份证号码
        SCBG_find_replace(SCBG, "{强制措施}", ExtraDef.catch(info).lstrip('涉嫌危险驾驶罪于'))  # 强制措施
        SCBG_find_replace(SCBG, "执勤民警", "温岭市公安局交警大队民警")  # 民警替换
        SCBG_find_replace(SCBG, "XXX", inquisitor)  # 检察官名字替换
        SCBG_find_replace(SCBG, "经温岭市公安司法鉴定中心鉴定", "经抽血检验")  # 经……抽血检验
        SCBG_find_replace(SCBG, "{标题姓名}", R_name)  # 标题姓名
        for i_page in range(14):  # 页码
            SCBG_find_replace("审查报告" + "（" + R_name + "）.docx", "{页码" + str(i_page + 1) + "}", pages[i_page])
        doc = Document(SCBG)  # 字体
        doc.styles['Normal'].font.name = '仿宋_GB2312'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
        doc.styles['Normal'].font.size = Pt(14)
        # doc.save(SCBG)  # ?
        # modify_specific_word(doc, R_name + "危险驾驶案审查报告", 22, '宋体')

    else:
        pages = ScbgScan.get_bookmarks_2(desktop_path + '/诉讼证据卷.pdf')  # 获取页码
        if lxjy_C == '相对不起诉':
            SCBG_find_replace(r"模板/dinglin/不诉审查报告.docx", "{基础信息}", info_new)  # 基础信息
            SCBG_find_replace(SCBG, "{姓名}", R_name)  # 姓名
            SCBG_find_replace(SCBG, "{犯罪事实}", fact.lstrip("经依法侦查查明："))  # 犯罪事实
            SCBG_find_replace(SCBG, "{量刑建议}", lxjy_C)  # 量刑建议
            SCBG_find_replace(SCBG, "{性别}", sex)  # 性别
            SCBG_find_replace(SCBG, "{年龄}", age)  # 年龄
            SCBG_find_replace(SCBG, "{身份证号码}", id_number)  # 身份证号码
            SCBG_find_replace(SCBG, "{前科}", qianke)  # 前科劣迹
            SCBG_find_replace(SCBG, "{酒精含量}", J_middle_word + 'mg/100ml')  # 酒精含量
            SCBG_find_replace(SCBG, "{强制措施}",
                              ExtraDef.catch(info).lstrip('涉嫌危险驾驶罪于'))  # 强制措施
            SCBG_find_replace(SCBG, "执勤民警", "温岭市公安局交警大队民警")  # 民警替换
            SCBG_find_replace(SCBG, "XXX", inquisitor)  # 检察官名字替换
            SCBG_find_replace(SCBG, "经温岭市公安司法鉴定中心鉴定", "经抽血检验")  # 经……抽血检验
            for i_page in range(18):  # 页码
                SCBG_find_replace("审查报告" + "（" + R_name + "）.docx", "{页码" + str(i_page + 1) + "}", pages[i_page])
            SCBG_find_replace(SCBG, "{标题姓名}", R_name)  # 标题姓名
            doc = Document(SCBG)
            doc.styles['Normal'].font.name = '仿宋_GB2312'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
            doc.styles['Normal'].font.size = Pt(14)
        else:
            SCBG_find_replace(r"模板/dinglin/起诉审查报告.docx", "{基础信息}", info_new)  # 基础信息
            SCBG_find_replace("审查报告" + "（" + R_name + "）.docx", "{姓名}", R_name)  # 姓名
            SCBG_find_replace("审查报告" + "（" + R_name + "）.docx", "{性别}", sex)  # 性别
            SCBG_find_replace("审查报告" + "（" + R_name + "）.docx", "{年龄}", age)  # 年龄
            SCBG_find_replace("审查报告" + "（" + R_name + "）.docx", "{犯罪事实}", fact.lstrip("经依法侦查查明："))  # 犯罪事实
            SCBG_find_replace("审查报告" + "（" + R_name + "）.docx", "{量刑建议}", lxjy_C)  # 量刑建议
            SCBG_find_replace("审查报告" + "（" + R_name + "）.docx", "{前科}", qianke)  # 前科劣迹
            SCBG_find_replace("审查报告" + "（" + R_name + "）.docx", "{强制措施}",
                              ExtraDef.catch(info).lstrip('涉嫌危险驾驶罪于'))  # 强制措施
            SCBG_find_replace("审查报告" + "（" + R_name + "）.docx", "{身份证号码}", id_number)  # 身份证号码
            SCBG_find_replace("审查报告" + "（" + R_name + "）.docx", "执勤民警", "温岭市公安局交警大队民警")  # 民警替换
            SCBG_find_replace("审查报告" + "（" + R_name + "）.docx", "XXX", inquisitor)  # 检察官名字替换
            SCBG_find_replace("审查报告" + "（" + R_name + "）.docx", "经温岭市公安司法鉴定中心鉴定", "经抽血检验")  # 经……抽血检验
            for i_page in range(18):  # 页码
                SCBG_find_replace("审查报告" + "（" + R_name + "）.docx", "{页码" + str(i_page + 1) + "}", pages[i_page])
            SCBG_find_replace(SCBG, "{标题姓名}", R_name)  # 标题姓名
            doc = Document("审查报告（" + R_name + "）.docx")
            doc.styles['Normal'].font.name = '仿宋_GB2312'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
            doc.styles['Normal'].font.size = Pt(14)

    search_text = info_new
    add_underline_to_text(doc, search_text)
    search_text2 = fact.lstrip("经依法侦查查明：")
    add_underline_to_text(doc, search_text2)
    search_text = lxjy_C
    add_underline_to_text(doc, search_text)
    doc.save(SCBG)

    '''移动到桌面'''

    # 移动文件
    def movefile(srcfile, dstpath):
        if not os.path.isfile(srcfile):
            print("%s not exist!" % (srcfile))
        else:
            fpath, fname = os.path.split(srcfile)  # 分离文件名和路径
            if not os.path.exists(dstpath):
                os.makedirs(dstpath)  # 创建路径
            shutil.move(srcfile, dstpath + fname)  # 移动文件
            # print("move %s -> %s" % (srcfile, dstpath + fname))

    if lxjy_C == "相对不起诉。":
        movefile("落实“三个规定”报告表（" + R_name + "）.docx", desktop_path + '/')
        movefile("不起诉决定书（" + R_name + "）.docx", desktop_path + '/')
        movefile("认罪认罚具结书（" + R_name + "）.docx", desktop_path + '/')
        movefile("讯问笔录（" + R_name + "）.docx", desktop_path + '/')
        movefile("审查报告（" + R_name + "）.docx", desktop_path + '/')
    else:
        movefile("起诉书（" + R_name + "）.docx", desktop_path + '/')
        movefile("落实“三个规定”报告表（" + R_name + "）.docx", desktop_path + '/')
        movefile("认罪认罚具结书（" + R_name + "）.docx", desktop_path + '/')
        movefile("公诉意见书（" + R_name + "）.docx", desktop_path + '/')
        movefile("讯问笔录（" + R_name + "）.docx", desktop_path + '/')
        movefile("审查报告（" + R_name + "）.docx", desktop_path + '/')

    # 删除output.txt
    os.remove('output.txt')

    # 最终字体调整
    path = desktop_path + '/审查报告（' + R_name + '）.docx'
    change_text(path, R_name + '危险驾驶案审查报告', '宋体', 22, True, RGBColor(0, 0, 0), WD_PARAGRAPH_ALIGNMENT.CENTER)  # 加粗
    change_text(path, '关于犯罪嫌疑人' + R_name, '方正小标宋简体', 22, True, RGBColor(0, 0, 0), WD_PARAGRAPH_ALIGNMENT.CENTER)  # 加粗
    # 打开文件夹
    try:
        os.startfile(desktop_path)
    except:
        print("打开文件夹失败")

    return print("文件已生成至桌面" + R_name + "文件夹中")

    # print(LX('/Users/xumuzhi/coding/pdfGet/起诉意见书.pdf'))         # Mac
    # print(LX(r'C:\Users\Administrator\Desktop\测试用起诉意见书\起诉意见书 (1).pdf'))        # windows


# print(SC('output.txt'))  # windows
